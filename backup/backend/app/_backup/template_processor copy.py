"""
═══════════════════════════════════════════════════════════════
TEMPLATE PROCESSOR - Procesador de Plantillas Word
═══════════════════════════════════════════════════════════════

PROPÓSITO:
Procesar plantillas Word (.docx) personalizadas con marcadores
y reemplazarlos con datos reales para generar documentos finales.

MARCADORES SOPORTADOS:
- {{cliente}}           → Nombre del cliente
- {{proyecto}}          → Nombre del proyecto
- {{numero_cotizacion}} → Número de cotización
- {{fecha}}             → Fecha actual
- {{descripcion}}       → Descripción del proyecto
- {{subtotal}}          → Subtotal
- {{igv}}               → IGV (18%)
- {{total}}             → Total
- {{items_tabla}}       → Tabla de items (se reemplaza con tabla real)
- {{logo}}              → Logo de empresa (se inserta imagen)

USO:
    from app.services.template_processor import template_processor
    
    datos = {
        "cliente": "Acme Corp",
        "numero": "COT-202510-0001",
        "items": [...]
    }
    
    archivo = template_processor.procesar_plantilla(
        ruta_plantilla="storage/plantillas/mi_plantilla.docx",
        datos_cotizacion=datos,
        logo_base64="data:image/png;base64,..."
    )

═══════════════════════════════════════════════════════════════
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime
import logging
import re
import base64
from io import BytesIO

logger = logging.getLogger(__name__)

class TemplateProcessor:
    """
    Procesador de plantillas Word con marcadores
    
    Características:
    - Reemplaza marcadores en texto
    - Inserta tablas dinámicamente
    - Agrega logos/imágenes
    - Mantiene formato original
    - Valida plantillas
    """
    
    def __init__(self, plantillas_dir: str = "backend/storage/plantillas"):
        """
        Inicializar procesador
        
        Args:
            plantillas_dir: Directorio donde se guardan plantillas
        """
        self.plantillas_dir = Path(plantillas_dir)
        self.plantillas_dir.mkdir(parents=True, exist_ok=True)
        
        # Marcadores válidos
        self.marcadores_validos = [
            'cliente', 'proyecto', 'numero_cotizacion', 'fecha',
            'descripcion', 'subtotal', 'igv', 'total', 'estado',
            'items_tabla', 'logo', 'observaciones', 'vigencia'
        ]
        
        logger.info(f"TemplateProcessor inicializado. Dir: {self.plantillas_dir}")
    
    # ════════════════════════════════════════════════════════
    # FUNCIÓN PRINCIPAL
    # ════════════════════════════════════════════════════════
    
    def procesar_plantilla(
        self,
        ruta_plantilla: str,
        datos_cotizacion: Dict[str, Any],
        ruta_salida: Optional[str] = None,
        logo_base64: Optional[str] = None
    ) -> str:
        """
        Procesar plantilla Word con datos de cotización/proyecto
        
        Args:
            ruta_plantilla: Ruta a la plantilla .docx
            datos_cotizacion: Datos para reemplazar marcadores
            ruta_salida: Ruta de salida (opcional)
            logo_base64: Logo en base64 (opcional)
        
        Returns:
            Ruta del documento generado
        
        Raises:
            Exception: Si hay error al procesar
        """
        
        try:
            logger.info(f"Procesando plantilla: {ruta_plantilla}")
            
            # Validar que existe la plantilla
            if not Path(ruta_plantilla).exists():
                raise FileNotFoundError(f"Plantilla no encontrada: {ruta_plantilla}")
            
            # Abrir plantilla
            doc = Document(ruta_plantilla)
            
            # Preparar datos con valores por defecto
            datos = self._preparar_datos(datos_cotizacion)
            
            # 1. Reemplazar marcadores en párrafos
            self._reemplazar_en_parrafos(doc, datos)
            
            # 2. Reemplazar marcadores en tablas
            self._reemplazar_en_tablas(doc, datos)
            
            # 3. Reemplazar marcadores en headers/footers
            self._reemplazar_en_headers_footers(doc, datos)
            
            # 4. Insertar tabla de items si hay marcador
            if '{{items_tabla}}' in self._obtener_todo_texto(doc):
                self._insertar_tabla_items(doc, datos_cotizacion.get('items', []))
            
            # 5. Insertar logo si existe
            if logo_base64 and '{{logo}}' in self._obtener_todo_texto(doc):
                self._insertar_logo(doc, logo_base64)
            
            # Definir ruta de salida
            if not ruta_salida:
                from app.core.config import settings
                numero = datos_cotizacion.get('numero', 'TEMP')
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nombre_archivo = f"documento_{numero}_{timestamp}.docx"
                ruta_salida = str(Path(settings.GENERATED_DIR) / nombre_archivo)
            
            # Crear directorio si no existe
            Path(ruta_salida).parent.mkdir(parents=True, exist_ok=True)
            
            # Guardar documento
            doc.save(ruta_salida)
            
            logger.info(f"✅ Documento generado: {ruta_salida}")
            
            return ruta_salida
            
        except Exception as e:
            logger.error(f"❌ Error al procesar plantilla: {str(e)}")
            raise Exception(f"Error al procesar plantilla: {str(e)}")
    
    # ════════════════════════════════════════════════════════
    # PREPARACIÓN DE DATOS
    # ════════════════════════════════════════════════════════
    
    def _preparar_datos(self, datos: Dict[str, Any]) -> Dict[str, str]:
        """
        Preparar datos para reemplazo en plantilla
        
        Convierte todos los valores a strings formateados
        
        Args:
            datos: Datos raw del proyecto/cotización
        
        Returns:
            Dict con marcadores y valores formateados
        """
        
        # Fecha actual formateada
        meses_es = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        hoy = datetime.now()
        fecha_formateada = f"{hoy.day} de {meses_es[hoy.month]} de {hoy.year}"
        
        # Calcular totales si no existen
        subtotal = float(datos.get('subtotal', 0))
        igv = float(datos.get('igv', 0))
        total = float(datos.get('total', 0))
        
        # Si no hay subtotal pero hay items, calcular
        if subtotal == 0 and datos.get('items'):
            subtotal = sum(
                float(item.get('cantidad', 0)) * float(item.get('precio_unitario', 0))
                for item in datos.get('items', [])
            )
            igv = subtotal * 0.18
            total = subtotal + igv
        
        # Crear diccionario de reemplazo
        datos_procesados = {
            '{{cliente}}': str(datos.get('cliente', 'N/A')),
            '{{proyecto}}': str(datos.get('proyecto', 'N/A')),
            '{{numero_cotizacion}}': str(datos.get('numero', 'N/A')),
            '{{fecha}}': fecha_formateada,
            '{{descripcion}}': str(datos.get('descripcion', '')),
            '{{subtotal}}': f"S/ {subtotal:,.2f}",
            '{{igv}}': f"S/ {igv:,.2f}",
            '{{total}}': f"S/ {total:,.2f}",
            '{{estado}}': str(datos.get('estado', 'borrador')).upper(),
            '{{observaciones}}': str(datos.get('observaciones', '')),
            '{{vigencia}}': str(datos.get('vigencia', '30 días')),
        }
        
        return datos_procesados
    
    # ════════════════════════════════════════════════════════
    # REEMPLAZO EN DIFERENTES PARTES DEL DOCUMENTO
    # ════════════════════════════════════════════════════════
    
    def _reemplazar_en_parrafos(self, doc: Document, datos: Dict[str, str]):
        """
        Reemplazar marcadores en todos los párrafos del documento
        
        Mantiene el formato original del texto
        """
        
        for parrafo in doc.paragraphs:
            self._reemplazar_en_parrafo(parrafo, datos)
    
    def _reemplazar_en_parrafo(self, parrafo, datos: Dict[str, str]):
        """
        Reemplazar marcadores en un párrafo específico
        
        Usa runs para mantener formato
        """
        
        # Buscar marcadores en el texto completo del párrafo
        texto_completo = parrafo.text
        
        for marcador, valor in datos.items():
            if marcador in texto_completo:
                # Reemplazar en cada run
                for run in parrafo.runs:
                    if marcador in run.text:
                        run.text = run.text.replace(marcador, valor)
    
    def _reemplazar_en_tablas(self, doc: Document, datos: Dict[str, str]):
        """
        Reemplazar marcadores en todas las tablas
        """
        
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        self._reemplazar_en_parrafo(parrafo, datos)
    
    def _reemplazar_en_headers_footers(self, doc: Document, datos: Dict[str, str]):
        """
        Reemplazar marcadores en encabezados y pies de página
        """
        
        for section in doc.sections:
            # Header
            for parrafo in section.header.paragraphs:
                self._reemplazar_en_parrafo(parrafo, datos)
            
            # Footer
            for parrafo in section.footer.paragraphs:
                self._reemplazar_en_parrafo(parrafo, datos)
    
    # ════════════════════════════════════════════════════════
    # INSERCIÓN DE TABLA DE ITEMS
    # ════════════════════════════════════════════════════════
    
    def _insertar_tabla_items(self, doc: Document, items: List[Dict]):
        """
        Busca el marcador {{items_tabla}} y lo reemplaza con una tabla real
        
        Args:
            doc: Documento Word
            items: Lista de items para la tabla
        """
        
        if not items:
            logger.warning("No hay items para insertar en la tabla")
            return
        
        # Buscar el marcador
        for i, parrafo in enumerate(doc.paragraphs):
            if '{{items_tabla}}' in parrafo.text:
                logger.info(f"Encontrado marcador items_tabla en párrafo {i}")
                
                # Limpiar el párrafo
                parrafo.clear()
                
                # Crear tabla después del párrafo
                # Columnas: Descripción, Cantidad, P.Unit, Total
                tabla = doc.add_table(rows=1, cols=4)
                tabla.style = 'Light Grid Accent 1'
                
                # Encabezados
                celdas_header = tabla.rows[0].cells
                celdas_header[0].text = 'Descripción'
                celdas_header[1].text = 'Cantidad'
                celdas_header[2].text = 'P. Unitario'
                celdas_header[3].text = 'Total'
                
                # Formatear encabezados
                for celda in celdas_header:
                    for p in celda.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    # Color de fondo (requiere XML)
                    celda._element.get_or_add_tcPr().append(
                        self._create_shading_element('8B0000')  # Rojo Tesla
                    )
                
                # Agregar items
                for item in items:
                    fila_cells = tabla.add_row().cells
                    
                    # Descripción
                    fila_cells[0].text = str(item.get('descripcion', ''))
                    
                    # Cantidad
                    cantidad = float(item.get('cantidad', 0))
                    fila_cells[1].text = f"{cantidad:.2f}"
                    fila_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Precio Unitario
                    precio_unit = float(item.get('precio_unitario', 0))
                    fila_cells[2].text = f"S/ {precio_unit:,.2f}"
                    fila_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                    # Total
                    total_item = cantidad * precio_unit
                    fila_cells[3].text = f"S/ {total_item:,.2f}"
                    fila_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                logger.info(f"✅ Tabla insertada con {len(items)} items")
                
                # Solo procesar el primer marcador encontrado
                break
    
    def _create_shading_element(self, color_hex: str):
        """
        Crear elemento de sombreado para celdas de tabla
        
        Args:
            color_hex: Color en hexadecimal (ej: '8B0000')
        
        Returns:
            Elemento XML de sombreado
        """
        from docx.oxml import parse_xml
        
        shading_xml = f'<w:shd {{{qn("w")}}} w:fill="{color_hex}"/>'
        return parse_xml(shading_xml)
    
    # ════════════════════════════════════════════════════════
    # INSERCIÓN DE LOGO
    # ════════════════════════════════════════════════════════
    
    def _insertar_logo(self, doc: Document, logo_base64: str):
        """
        Busca el marcador {{logo}} y lo reemplaza con la imagen
        
        Args:
            doc: Documento Word
            logo_base64: Logo en formato base64
        """
        
        try:
            # Decodificar base64
            if ',' in logo_base64:
                logo_base64 = logo_base64.split(',')[1]
            
            imagen_bytes = base64.b64decode(logo_base64)
            imagen_io = BytesIO(imagen_bytes)
            
            # Buscar marcador
            logo_insertado = False
            
            for parrafo in doc.paragraphs:
                if '{{logo}}' in parrafo.text:
                    # Limpiar párrafo
                    parrafo.clear()
                    
                    # Insertar imagen
                    run = parrafo.add_run()
                    run.add_picture(imagen_io, width=Inches(2))
                    
                    # Centrar
                    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    logo_insertado = True
                    logger.info("✅ Logo insertado en marcador")
                    break
            
            # Si no hay marcador, insertar al principio
            if not logo_insertado:
                parrafo_logo = doc.paragraphs[0].insert_paragraph_before()
                run = parrafo_logo.add_run()
                run.add_picture(imagen_io, width=Inches(2))
                parrafo_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.info("✅ Logo insertado al inicio del documento")
            
        except Exception as e:
            logger.warning(f"⚠️ No se pudo insertar logo: {str(e)}")
    
    # ════════════════════════════════════════════════════════
    # UTILIDADES
    # ════════════════════════════════════════════════════════
    
    def _obtener_todo_texto(self, doc: Document) -> str:
        """
        Obtener todo el texto del documento (para búsqueda de marcadores)
        
        Returns:
            Todo el texto concatenado
        """
        
        texto_completo = []
        
        # Párrafos
        for p in doc.paragraphs:
            texto_completo.append(p.text)
        
        # Tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        texto_completo.append(p.text)
        
        return ' '.join(texto_completo)
    
    def validar_plantilla(self, ruta_plantilla: str) -> Dict[str, Any]:
        """
        Validar una plantilla y extraer marcadores encontrados
        
        Args:
            ruta_plantilla: Ruta a la plantilla
        
        Returns:
            Dict con información de validación
        """
        
        try:
            doc = Document(ruta_plantilla)
            texto_completo = self._obtener_todo_texto(doc)
            
            # Buscar marcadores
            marcadores_encontrados = re.findall(r'\{\{(\w+)\}\}', texto_completo)
            marcadores_encontrados = list(set(marcadores_encontrados))
            
            # Verificar si son válidos
            marcadores_invalidos = [
                m for m in marcadores_encontrados 
                if m not in self.marcadores_validos
            ]
            
            es_valida = len(marcadores_invalidos) == 0
            
            return {
                "valida": es_valida,
                "marcadores_encontrados": marcadores_encontrados,
                "marcadores_invalidos": marcadores_invalidos,
                "marcadores_validos": self.marcadores_validos,
                "total_parrafos": len(doc.paragraphs),
                "total_tablas": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"Error al validar plantilla: {str(e)}")
            return {
                "valida": False,
                "error": str(e)
            }
    
    def listar_plantillas(self) -> List[Dict[str, Any]]:
        """
        Listar todas las plantillas disponibles
        
        Returns:
            Lista de plantillas con metadata
        """
        
        plantillas = []
        
        for archivo in self.plantillas_dir.glob("*.docx"):
            if not archivo.name.startswith('~'):  # Ignorar temporales
                
                # Validar plantilla
                validacion = self.validar_plantilla(str(archivo))
                
                plantillas.append({
                    "nombre": archivo.name,
                    "ruta": str(archivo),
                    "tamano": archivo.stat().st_size,
                    "fecha_modificacion": datetime.fromtimestamp(
                        archivo.stat().st_mtime
                    ).isoformat(),
                    "valida": validacion.get("valida", False),
                    "marcadores": validacion.get("marcadores_encontrados", [])
                })
        
        return plantillas
    
    def eliminar_plantilla(self, nombre_plantilla: str) -> bool:
        """
        Eliminar una plantilla
        
        Args:
            nombre_plantilla: Nombre del archivo
        
        Returns:
            True si se eliminó exitosamente
        """
        
        try:
            ruta = self.plantillas_dir / nombre_plantilla
            
            if ruta.exists():
                ruta.unlink()
                logger.info(f"Plantilla eliminada: {nombre_plantilla}")
                return True
            else:
                logger.warning(f"Plantilla no encontrada: {nombre_plantilla}")
                return False
                
        except Exception as e:
            logger.error(f"Error al eliminar plantilla: {str(e)}")
            return False

# ════════════════════════════════════════════════════════
# INSTANCIA GLOBAL
# ════════════════════════════════════════════════════════

template_processor = TemplateProcessor()