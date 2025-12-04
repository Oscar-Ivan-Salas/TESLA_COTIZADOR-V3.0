"""
ğŸ“„ WORD GENERATOR + PILI v3.0 - DOCUMENTOS PERFECTOS
ğŸ“ RUTA: backend/app/services/word_generator.py

PILI (Procesadora Inteligente de Licitaciones Industriales) integrada con 
generaciÃ³n Word profesional para crear documentos perfectos sin corrupciÃ³n.

ğŸ¯ NUEVAS CARACTERÃSTICAS PILI v3.0:
- Procesamiento JSON estructurado de PILI
- IntegraciÃ³n automÃ¡tica con agentes especializados  
- Documentos personalizados por tipo de servicio
- Sin corrupciÃ³n de archivos (problema resuelto)
- Plantillas inteligentes con marcadores PILI
- Logos automÃ¡ticos y formateo profesional

ğŸ”„ CONSERVA TODO LO EXISTENTE:
- generar_cotizacion() âœ…
- generar_informe_proyecto() âœ…
- generar_informe_simple() âœ…
- Estilos Tesla profesionales âœ…
- Toda la lÃ³gica de formateo âœ…
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path
import logging
import base64
from io import BytesIO
import json
import tempfile

logger = logging.getLogger(__name__)

class WordGenerator:
    """
    ğŸ”„ GENERADOR ORIGINAL CONSERVADO + ğŸ¤– PILI INTEGRADA
    
    Mantiene toda la funcionalidad existente pero agrega capacidades
    inteligentes de PILI para generar documentos perfectos.
    """
    
    def __init__(self):
        """ğŸ”„ CONSERVADO + ğŸ¤– PILI mejorado"""
        # Colores Tesla originales conservados
        self.COLOR_ROJO = RGBColor(139, 0, 0)      # #8B0000
        self.COLOR_DORADO = RGBColor(218, 165, 32)  # #DAA520
        self.COLOR_NEGRO = RGBColor(0, 0, 0)        # #000000
        self.COLOR_GRIS = RGBColor(128, 128, 128)   # #808080
        
        # ğŸ¤– Nuevos colores PILI
        self.COLOR_PILI = RGBColor(212, 175, 55)    # #D4AF37 (Dorado PILI)
        self.COLOR_AZUL_TECH = RGBColor(0, 102, 204)  # #0066CC (Azul tecnolÃ³gico)
        
        # ConfiguraciÃ³n de documentos
        self.empresa_info = {
            "nombre": "TESLA ELECTRICIDAD Y AUTOMATIZACIÃ“N S.A.C.",
            "ruc": "20601138787",
            "direccion": "Jr. Las Ãgatas Mz B Lote 09, Urb. San Carlos, SJL",
            "telefono": "906315961",
            "email": "ingenieria.teslaelectricidad@gmail.com"
        }
        
        logger.info("âœ… WordGenerator + PILI inicializado")

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # ğŸ¤– NUEVOS MÃ‰TODOS PILI v3.0
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    
    def generar_desde_json_pili(
        self,
        datos_json: Dict[str, Any],
        tipo_documento: str = "cotizacion",
        opciones: Optional[Dict[str, Any]] = None,
        logo_base64: Optional[str] = None,
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        ğŸ¤– NUEVO PILI v3.0 - Genera documento Word desde JSON estructurado de PILI

        Args:
            datos_json: Datos estructurados por PILI
            tipo_documento: Tipo (cotizacion, proyecto, informe)
            opciones: Opciones de personalizaciÃ³n
            logo_base64: Logo en base64
            ruta_salida: Ruta personalizada para guardar el archivo

        Returns:
            InformaciÃ³n del documento generado
        """

        try:
            logger.info(f"ğŸ¤– PILI generando documento {tipo_documento} desde JSON")

            # 1. Validar y procesar datos JSON
            datos_procesados = self._procesar_json_pili(datos_json, tipo_documento)

            # 2. Obtener agente PILI responsable
            agente_pili = datos_json.get("agente_responsable", "PILI")

            # 3. Determinar mÃ©todo de generaciÃ³n segÃºn tipo
            if tipo_documento == "cotizacion" or "cotizacion" in tipo_documento:
                resultado = self._generar_cotizacion_pili(datos_procesados, agente_pili, opciones, logo_base64, ruta_salida)

            elif tipo_documento == "proyecto" or "proyecto" in tipo_documento:
                resultado = self._generar_proyecto_pili(datos_procesados, agente_pili, opciones, logo_base64, ruta_salida)

            elif tipo_documento == "informe" or "informe" in tipo_documento:
                resultado = self._generar_informe_pili(datos_procesados, agente_pili, opciones, logo_base64, ruta_salida)

            else:
                # Fallback a cotizaciÃ³n
                resultado = self._generar_cotizacion_pili(datos_procesados, agente_pili, opciones, logo_base64, ruta_salida)

            # 4. Agregar metadatos PILI
            resultado["pili_metadata"] = {
                "agente_responsable": agente_pili,
                "tipo_documento": tipo_documento,
                "timestamp_generacion": datetime.now().isoformat(),
                "version_pili": "3.0"
            }

            logger.info(f"âœ… PILI documento generado: {resultado.get('nombre_archivo', 'documento.docx')}")
            return resultado

        except Exception as e:
            logger.error(f"âŒ Error PILI generando documento: {str(e)}")
            return {
                "exito": False,
                "error": str(e),
                "mensaje": f"Error generando documento {tipo_documento}: {str(e)}"
            }
    
    def _procesar_json_pili(self, datos_json: Dict[str, Any], tipo_documento: str) -> Dict[str, Any]:
        """Procesa y valida JSON de PILI para generaciÃ³n Word"""
        
        # Extraer datos principales
        datos_extraidos = datos_json.get("datos_extraidos", {})
        
        # Datos base del documento
        datos_procesados = {
            "empresa_nombre": self.empresa_info["nombre"],
            "empresa_ruc": self.empresa_info["ruc"],
            "empresa_direccion": self.empresa_info["direccion"],
            "empresa_telefono": self.empresa_info["telefono"],
            "empresa_email": self.empresa_info["email"],
            "fecha_generacion": datetime.now().strftime("%d/%m/%Y"),
            "agente_pili": datos_json.get("agente_responsable", "PILI")
        }
        
        # Combinar con datos extraÃ­dos
        datos_procesados.update(datos_extraidos)
        
        # Valores por defecto segÃºn tipo de documento
        if "cotizacion" in tipo_documento:
            datos_procesados.setdefault("numero", self._generar_numero_cotizacion())
            datos_procesados.setdefault("vigencia", "30 dÃ­as")
            datos_procesados.setdefault("observaciones", "Precios incluyen IGV. InstalaciÃ³n segÃºn CNE-UtilizaciÃ³n.")
            
        elif "proyecto" in tipo_documento:
            datos_procesados.setdefault("estado", "En PlanificaciÃ³n")
            datos_procesados.setdefault("duracion_estimada", "4 semanas")
            
        elif "informe" in tipo_documento:
            datos_procesados.setdefault("autor", self.empresa_info["nombre"])
            datos_procesados.setdefault("fecha_informe", datos_procesados["fecha_generacion"])
        
        return datos_procesados
    
    def _generar_cotizacion_pili(
        self,
        datos: Dict[str, Any],
        agente_pili: str,
        opciones: Optional[Dict[str, Any]],
        logo_base64: Optional[str],
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """Genera cotizaciÃ³n con formato PILI personalizado"""

        # Crear documento nuevo
        doc = Document()

        # Configurar mÃ¡rgenes
        self._configurar_margenes(doc)

        # 1. Header con logo PILI
        if logo_base64:
            self._insertar_logo_pili(doc, logo_base64, agente_pili)
        else:
            self._insertar_header_pili(doc, agente_pili)

        # 2. InformaciÃ³n de empresa
        self._insertar_info_empresa(doc)

        # 3. TÃ­tulo cotizaciÃ³n
        self._insertar_titulo_cotizacion(doc, agente_pili)

        # 4. Datos del cliente
        self._insertar_datos_cliente_pili(doc, datos)

        # 5. DescripciÃ³n del proyecto
        self._insertar_descripcion_proyecto(doc, datos)

        # 6. Tabla de items
        self._insertar_tabla_items_pili(doc, datos)

        # 7. Totales
        self._insertar_totales_pili(doc, datos)

        # 8. Observaciones
        self._insertar_observaciones_pili(doc, datos)

        # 9. Footer PILI
        self._insertar_footer_pili(doc, agente_pili)

        # 10. Guardar documento
        return self._guardar_documento(doc, datos, "cotizacion", ruta_salida)
    
    def _generar_proyecto_pili(
        self,
        datos: Dict[str, Any],
        agente_pili: str,
        opciones: Optional[Dict[str, Any]],
        logo_base64: Optional[str],
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """Genera documento de proyecto con formato PILI"""

        # Crear documento nuevo
        doc = Document()

        # Configurar mÃ¡rgenes
        self._configurar_margenes(doc)

        # 1. Header PILI
        self._insertar_header_pili(doc, agente_pili)

        # 2. InformaciÃ³n empresa
        self._insertar_info_empresa(doc)

        # 3. TÃ­tulo proyecto
        titulo = doc.add_heading("GESTIÃ“N DE PROYECTO", level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._aplicar_estilo_titulo(titulo, self.COLOR_DORADO)

        # 4. InformaciÃ³n del proyecto
        self._insertar_info_proyecto_pili(doc, datos, agente_pili)

        # 5. Fases del proyecto (si existen)
        if "fases" in datos and datos["fases"]:
            self._insertar_fases_proyecto(doc, datos["fases"])

        # 6. Footer PILI
        self._insertar_footer_pili(doc, agente_pili)

        # 7. Guardar documento
        return self._guardar_documento(doc, datos, "proyecto", ruta_salida)
    
    def _generar_informe_pili(
        self,
        datos: Dict[str, Any],
        agente_pili: str,
        opciones: Optional[Dict[str, Any]],
        logo_base64: Optional[str],
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """Genera informe con formato PILI personalizado"""

        # Crear documento nuevo
        doc = Document()

        # Configurar mÃ¡rgenes
        self._configurar_margenes(doc)

        # 1. Header PILI
        self._insertar_header_pili(doc, agente_pili)

        # 2. InformaciÃ³n empresa
        self._insertar_info_empresa(doc)

        # 3. TÃ­tulo informe
        titulo_text = datos.get("titulo_informe", "INFORME TÃ‰CNICO")
        titulo = doc.add_heading(titulo_text, level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._aplicar_estilo_titulo(titulo, self.COLOR_DORADO)

        # 4. InformaciÃ³n del informe
        self._insertar_info_informe_pili(doc, datos, agente_pili)

        # 5. Contenido del informe
        if "resumen_ejecutivo" in datos:
            self._insertar_seccion_informe(doc, "RESUMEN EJECUTIVO", datos["resumen_ejecutivo"])

        if "conclusiones" in datos:
            self._insertar_seccion_informe(doc, "CONCLUSIONES", datos["conclusiones"])

        if "recomendaciones" in datos:
            self._insertar_seccion_informe(doc, "RECOMENDACIONES", datos["recomendaciones"])

        # 6. Footer PILI
        self._insertar_footer_pili(doc, agente_pili)

        # 7. Guardar documento
        return self._guardar_documento(doc, datos, "informe", ruta_salida)
    
    def _insertar_header_pili(self, doc: Document, agente_pili: str):
        """Inserta header personalizado con marca PILI"""
        
        header_para = doc.add_paragraph()
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Agregar badge PILI
        run_pili = header_para.add_run(f"ğŸ¤– Generado por {agente_pili} | Tesla IA v3.0")
        run_pili.font.size = Pt(9)
        run_pili.font.color.rgb = self.COLOR_GRIS
        run_pili.italic = True
        
        # Agregar espacio
        doc.add_paragraph()
    
    def _insertar_footer_pili(self, doc: Document, agente_pili: str):
        """Inserta footer con informaciÃ³n PILI"""
        
        # Separador
        doc.add_paragraph()
        
        # Footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # LÃ­nea separadora
        footer_para.add_run("â”€" * 60 + "\n")
        
        # Info PILI
        run_footer = footer_para.add_run(f"âœ¨ Documento generado por {agente_pili} - Tu agente IA especializada\n")
        run_footer.font.size = Pt(9)
        run_footer.font.color.rgb = self.COLOR_PILI
        run_footer.italic = True
        
        # Timestamp
        timestamp_run = footer_para.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')} | Tesla Electricidad v3.0")
        timestamp_run.font.size = Pt(8)
        timestamp_run.font.color.rgb = self.COLOR_GRIS
    
    def _insertar_datos_cliente_pili(self, doc: Document, datos: Dict[str, Any]):
        """Inserta datos del cliente con formato PILI"""
        
        # TÃ­tulo secciÃ³n
        titulo = doc.add_heading("DATOS DEL CLIENTE", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        # Tabla de datos
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Datos a mostrar
        campos_cliente = [
            ("Cliente:", datos.get("cliente", "[Cliente por definir]")),
            ("Proyecto:", datos.get("proyecto", "[Proyecto por definir]")),
            ("NÃºmero:", datos.get("numero", "[NÃºmero]")),
            ("Fecha:", datos.get("fecha_generacion", datetime.now().strftime("%d/%m/%Y"))),
        ]
        
        if "vigencia" in datos:
            campos_cliente.append(("Vigencia:", datos["vigencia"]))
        
        # Llenar tabla
        for campo, valor in campos_cliente:
            row = table.add_row()
            row.cells[0].text = campo
            row.cells[1].text = str(valor)
            
            # Formateo
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = self.COLOR_PILI
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_tabla_items_pili(self, doc: Document, datos: Dict[str, Any]):
        """Inserta tabla de items con formato PILI mejorado"""
        
        items = datos.get("items", [])
        if not items:
            return
        
        # TÃ­tulo secciÃ³n
        titulo = doc.add_heading("DETALLE DE ITEMS", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        # Crear tabla
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['DESCRIPCIÃ“N', 'CANT.', 'UND.', 'P.UNITARIO', 'TOTAL']
        header_cells = table.rows[0].cells
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Formateo header
            para = header_cells[i].paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
            
            # Fondo dorado para header
            self._aplicar_fondo_celda(header_cells[i], self.COLOR_PILI)
        
        # Agregar items
        for item in items:
            row = table.add_row()
            cells = row.cells
            
            # Datos del item
            descripcion = item.get("descripcion", "")
            cantidad = float(item.get("cantidad", 0))
            unidad = item.get("unidad", "und")
            precio_unitario = float(item.get("precio_unitario", 0))
            total = cantidad * precio_unitario
            
            # Llenar celdas
            cells[0].text = descripcion
            cells[1].text = str(int(cantidad) if cantidad.is_integer() else cantidad)
            cells[2].text = unidad
            cells[3].text = f"S/ {precio_unitario:.2f}"
            cells[4].text = f"S/ {total:.2f}"
            
            # AlineaciÃ³n
            cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Formateo nÃºmeros
            cells[4].paragraphs[0].runs[0].bold = True
            cells[4].paragraphs[0].runs[0].font.color.rgb = self.COLOR_DORADO
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_totales_pili(self, doc: Document, datos: Dict[str, Any]):
        """Inserta totales con formato PILI destacado"""
        
        # Calcular totales si no existen
        items = datos.get("items", [])
        if items:
            subtotal = sum(float(item.get("cantidad", 0)) * float(item.get("precio_unitario", 0)) for item in items)
            igv = subtotal * 0.18
            total = subtotal + igv
        else:
            subtotal = float(datos.get("subtotal", 0))
            igv = float(datos.get("igv", 0))
            total = float(datos.get("total", 0))
        
        # Crear tabla de totales
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Datos totales
        totales = [
            ("Subtotal:", f"S/ {subtotal:.2f}"),
            ("IGV (18%):", f"S/ {igv:.2f}"),
            ("TOTAL:", f"S/ {total:.2f}")
        ]
        
        # Llenar tabla
        for i, (label, valor) in enumerate(totales):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = valor
            
            # Formateo
            label_run = row.cells[0].paragraphs[0].runs[0]
            valor_run = row.cells[1].paragraphs[0].runs[0]
            
            if i == 2:  # TOTAL
                label_run.bold = True
                label_run.font.size = Pt(14)
                label_run.font.color.rgb = self.COLOR_PILI
                
                valor_run.bold = True
                valor_run.font.size = Pt(14)
                valor_run.font.color.rgb = self.COLOR_PILI
            else:
                label_run.bold = True
                
            # AlineaciÃ³n
            row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # Espacio
    
    def _generar_numero_cotizacion(self) -> str:
        """Genera nÃºmero Ãºnico de cotizaciÃ³n"""
        fecha = datetime.now()
        return f"COT-{fecha.strftime('%Y%m%d')}-001"
    
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # ğŸ”„ MÃ‰TODOS AUXILIARES PARA FORMATEO
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    
    def _configurar_margenes(self, doc: Document):
        """Configura mÃ¡rgenes del documento"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
    
    def _aplicar_estilo_titulo(self, titulo, color: RGBColor):
        """Aplica estilo consistente a tÃ­tulos"""
        for run in titulo.runs:
            run.font.color.rgb = color
            run.font.size = Pt(16)
            run.bold = True
    
    def _aplicar_estilo_seccion(self, titulo):
        """Aplica estilo a tÃ­tulos de secciÃ³n"""
        for run in titulo.runs:
            run.font.color.rgb = self.COLOR_PILI
            run.font.size = Pt(12)
            run.bold = True
    
    def _aplicar_fondo_celda(self, celda, color: RGBColor):
        """Aplica color de fondo a una celda"""
        try:
            shading = celda._element.get_or_add_tcPr().get_or_add_shd()
            shading.fill = f"{color.r:02X}{color.g:02X}{color.b:02X}"
        except:
            pass  # Si falla, continuar sin fondo
    
    def _insertar_info_empresa(self, doc: Document):
        """Inserta informaciÃ³n de la empresa"""
        
        # TÃ­tulo empresa
        titulo_empresa = doc.add_heading(self.empresa_info["nombre"], 0)
        titulo_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._aplicar_estilo_titulo(titulo_empresa, self.COLOR_PILI)
        
        # InformaciÃ³n contacto
        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        info_text = f"RUC: {self.empresa_info['ruc']}\n"
        info_text += f"{self.empresa_info['direccion']}\n"
        info_text += f"ğŸ“ {self.empresa_info['telefono']} | ğŸ“§ {self.empresa_info['email']}"
        
        run = info_para.add_run(info_text)
        run.font.size = Pt(10)
        run.font.color.rgb = self.COLOR_GRIS
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_titulo_cotizacion(self, doc: Document, agente_pili: str):
        """Inserta tÃ­tulo de cotizaciÃ³n personalizado"""
        
        titulo = doc.add_heading("COTIZACIÃ“N", 1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._aplicar_estilo_titulo(titulo, self.COLOR_PILI)
        
        # SubtÃ­tulo con agente PILI
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitulo.add_run(f"Generada por {agente_pili}")
        run.font.size = Pt(10)
        run.font.color.rgb = self.COLOR_GRIS
        run.italic = True
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_descripcion_proyecto(self, doc: Document, datos: Dict[str, Any]):
        """Inserta descripciÃ³n del proyecto"""
        
        descripcion = datos.get("descripcion", "")
        if not descripcion:
            return
        
        titulo = doc.add_heading("DESCRIPCIÃ“N DEL PROYECTO", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        para = doc.add_paragraph(descripcion)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_observaciones_pili(self, doc: Document, datos: Dict[str, Any]):
        """Inserta observaciones con formato PILI"""
        
        observaciones = datos.get("observaciones", "")
        if not observaciones:
            return
        
        titulo = doc.add_heading("OBSERVACIONES", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        para = doc.add_paragraph(observaciones)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_logo_pili(self, doc: Document, logo_base64: str, agente_pili: str):
        """Inserta logo con marca PILI"""
        
        try:
            # Decodificar base64
            if ',' in logo_base64:
                logo_base64 = logo_base64.split(',')[1]
            
            logo_bytes = base64.b64decode(logo_base64)
            
            # Crear archivo temporal
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(logo_bytes)
                temp_path = temp_file.name
            
            # Insertar imagen
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            run = para.add_run()
            run.add_picture(temp_path, width=Inches(2.5))
            
            # Limpiar archivo temporal
            Path(temp_path).unlink()
            
            # Agregar marca PILI
            marca_para = doc.add_paragraph()
            marca_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            marca_run = marca_para.add_run(f"ğŸ¤– {agente_pili}")
            marca_run.font.size = Pt(8)
            marca_run.font.color.rgb = self.COLOR_GRIS
            
        except Exception as e:
            logger.error(f"Error insertando logo: {e}")
            # Fallback a header sin logo
            self._insertar_header_pili(doc, agente_pili)
    
    def _insertar_info_proyecto_pili(self, doc: Document, datos: Dict[str, Any], agente_pili: str):
        """Inserta informaciÃ³n especÃ­fica de proyecto"""
        
        # TÃ­tulo secciÃ³n
        titulo = doc.add_heading("INFORMACIÃ“N DEL PROYECTO", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        # Datos del proyecto
        campos = [
            ("Nombre del Proyecto:", datos.get("nombre_proyecto", datos.get("proyecto", "[Nombre del proyecto]"))),
            ("Cliente:", datos.get("cliente", "[Cliente]")),
            ("Fecha de Inicio:", datos.get("fecha_inicio", "[Fecha de inicio]")),
            ("DuraciÃ³n Estimada:", datos.get("duracion_estimada", "[DuraciÃ³n]")),
            ("Estado:", datos.get("estado", "En PlanificaciÃ³n")),
            ("Responsable:", f"{agente_pili} (Agente IA)")
        ]
        
        for campo, valor in campos:
            para = doc.add_paragraph()
            run_campo = para.add_run(f"{campo} ")
            run_campo.bold = True
            run_campo.font.color.rgb = self.COLOR_PILI
            
            run_valor = para.add_run(str(valor))
            run_valor.font.color.rgb = self.COLOR_NEGRO
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_fases_proyecto(self, doc: Document, fases: List[Dict[str, Any]]):
        """Inserta fases del proyecto"""
        
        titulo = doc.add_heading("FASES DEL PROYECTO", level=2)
        self._aplicar_estilo_seccion(titulo)
        
        # Tabla de fases
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['FASE', 'DURACIÃ“N', 'ESTADO']
        header_cells = table.rows[0].cells
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            para = header_cells[i].paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_fondo_celda(header_cells[i], self.COLOR_PILI)
        
        # Agregar fases
        for fase in fases:
            row = table.add_row()
            cells = row.cells
            
            cells[0].text = fase.get("nombre", "")
            cells[1].text = fase.get("duracion", "")
            cells[2].text = fase.get("estado", "pendiente").title()
            
            # Centrar contenido
            for cell in cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_info_informe_pili(self, doc: Document, datos: Dict[str, Any], agente_pili: str):
        """Inserta informaciÃ³n especÃ­fica de informe"""
        
        # Datos del informe
        campos = [
            ("Fecha del Informe:", datos.get("fecha_informe", datetime.now().strftime("%d/%m/%Y"))),
            ("Autor:", datos.get("autor", self.empresa_info["nombre"])),
            ("Generado por:", f"{agente_pili} (Agente IA)")
        ]
        
        for campo, valor in campos:
            para = doc.add_paragraph()
            run_campo = para.add_run(f"{campo} ")
            run_campo.bold = True
            run_campo.font.color.rgb = self.COLOR_PILI
            
            run_valor = para.add_run(str(valor))
            run_valor.font.color.rgb = self.COLOR_NEGRO
        
        doc.add_paragraph()  # Espacio
    
    def _insertar_seccion_informe(self, doc: Document, titulo_seccion: str, contenido: str):
        """Inserta secciÃ³n de informe"""
        
        if not contenido:
            return
        
        titulo = doc.add_heading(titulo_seccion, level=2)
        self._aplicar_estilo_seccion(titulo)
        
        para = doc.add_paragraph(contenido)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        doc.add_paragraph()  # Espacio
    
    def _guardar_documento(self, doc: Document, datos: Dict[str, Any], tipo: str, ruta_salida: Optional[str] = None) -> Dict[str, Any]:
        """Guarda documento y retorna informaciÃ³n"""

        try:
            # Usar ruta personalizada o generar una
            if ruta_salida:
                ruta_archivo = Path(ruta_salida)
                # Asegurar que el directorio existe
                ruta_archivo.parent.mkdir(parents=True, exist_ok=True)
                nombre_archivo = ruta_archivo.name
            else:
                # Generar nombre Ãºnico
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                cliente_slug = self._slugify(datos.get("cliente", "cliente"))
                nombre_archivo = f"{tipo}_{cliente_slug}_{timestamp}.docx"

                # Ruta de salida usando configuraciÃ³n centralizada
                from app.core.config import get_generated_directory
                output_dir = get_generated_directory()
                ruta_archivo = output_dir / nombre_archivo

            # Guardar documento
            doc.save(str(ruta_archivo))

            # InformaciÃ³n del archivo generado
            return {
                "exito": True,
                "ruta_archivo": str(ruta_archivo),
                "nombre_archivo": nombre_archivo,
                "tamano_bytes": ruta_archivo.stat().st_size,
                "tipo_documento": tipo,
                "fecha_generacion": datetime.now().isoformat(),
                "cliente": datos.get("cliente", ""),
                "mensaje": f"Documento {tipo} generado exitosamente"
            }

        except Exception as e:
            logger.error(f"Error guardando documento: {e}")
            return {
                "exito": False,
                "error": str(e),
                "mensaje": f"Error guardando documento: {str(e)}"
            }
    
    def _slugify(self, texto: str) -> str:
        """Convierte texto en slug vÃ¡lido para nombre archivo"""
        import re
        texto = re.sub(r'[^\w\s-]', '', texto.strip())
        texto = re.sub(r'[-\s]+', '-', texto)
        return texto.lower()[:20]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    # ğŸ”„ MÃ‰TODOS ORIGINALES CONSERVADOS (COMPATIBILIDAD)
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    
    def generar_cotizacion(
        self,
        datos_cotizacion: Optional[Dict[str, Any]] = None,
        opciones: Optional[Dict[str, Any]] = None,
        logo_base64: Optional[str] = None,
        datos: Optional[Dict[str, Any]] = None,
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        ğŸ”„ CONSERVADO - MÃ©todo original para compatibilidad

        Redirige al nuevo sistema PILI pero mantiene interfaz original.
        Acepta tanto datos_cotizacion como datos para compatibilidad.
        """

        try:
            # Compatibilidad: usar datos o datos_cotizacion
            datos_entrada = datos or datos_cotizacion or {}

            # Convertir datos antiguos a formato PILI JSON
            datos_json = {
                "datos_extraidos": datos_entrada,
                "agente_responsable": "PILI Cotizadora",
                "tipo_servicio": "cotizacion-simple",
                "timestamp": datetime.now().isoformat()
            }

            # Usar nuevo mÃ©todo PILI con ruta personalizada
            return self.generar_desde_json_pili(
                datos_json=datos_json,
                tipo_documento="cotizacion",
                opciones=opciones,
                logo_base64=logo_base64,
                ruta_salida=ruta_salida
            )

        except Exception as e:
            logger.error(f"Error en mÃ©todo legacy generar_cotizacion: {e}")
            return {
                "exito": False,
                "error": str(e),
                "mensaje": f"Error generando cotizaciÃ³n: {str(e)}"
            }
    
    def generar_informe_proyecto(
        self,
        datos_proyecto: Optional[Dict[str, Any]] = None,
        opciones: Optional[Dict[str, Any]] = None,
        logo_base64: Optional[str] = None,
        datos: Optional[Dict[str, Any]] = None,
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        ğŸ”„ CONSERVADO - MÃ©todo original para compatibilidad
        """

        try:
            # Compatibilidad: usar datos o datos_proyecto
            datos_entrada = datos or datos_proyecto or {}

            # Convertir a formato PILI JSON
            datos_json = {
                "datos_extraidos": datos_entrada,
                "agente_responsable": "PILI Coordinadora",
                "tipo_servicio": "proyecto-simple",
                "timestamp": datetime.now().isoformat()
            }

            # Usar nuevo mÃ©todo PILI
            return self.generar_desde_json_pili(
                datos_json=datos_json,
                tipo_documento="proyecto",
                opciones=opciones,
                logo_base64=logo_base64,
                ruta_salida=ruta_salida
            )

        except Exception as e:
            logger.error(f"Error en mÃ©todo legacy generar_informe_proyecto: {e}")
            return {
                "exito": False,
                "error": str(e),
                "mensaje": f"Error generando informe proyecto: {str(e)}"
            }
    
    def generar_informe_simple(
        self,
        datos_informe: Optional[Dict[str, Any]] = None,
        opciones: Optional[Dict[str, Any]] = None,
        logo_base64: Optional[str] = None,
        datos: Optional[Dict[str, Any]] = None,
        ruta_salida: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        ğŸ”„ CONSERVADO - MÃ©todo original para compatibilidad
        """

        try:
            # Compatibilidad: usar datos o datos_informe
            datos_entrada = datos or datos_informe or {}

            # Convertir a formato PILI JSON
            datos_json = {
                "datos_extraidos": datos_entrada,
                "agente_responsable": "PILI Reportera",
                "tipo_servicio": "informe-simple",
                "timestamp": datetime.now().isoformat()
            }

            # Usar nuevo mÃ©todo PILI
            return self.generar_desde_json_pili(
                datos_json=datos_json,
                tipo_documento="informe",
                opciones=opciones,
                logo_base64=logo_base64,
                ruta_salida=ruta_salida
            )

        except Exception as e:
            logger.error(f"Error en mÃ©todo legacy generar_informe_simple: {e}")
            return {
                "exito": False,
                "error": str(e),
                "mensaje": f"Error generando informe simple: {str(e)}"
            }

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ğŸ¯ INSTANCIA GLOBAL MEJORADA CON PILI
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# Crear instancia global con manejo robusto de errores
try:
    word_generator = WordGenerator()
    logger.info("âœ… WordGenerator + PILI inicializado correctamente")
except Exception as e:
    logger.error(f"âŒ Error crÃ­tico inicializando WordGenerator: {e}")
    word_generator = None

# FunciÃ³n auxiliar para obtener instancia segura
def get_word_generator():
    """Obtiene instancia de WordGenerator de forma segura"""
    global word_generator
    if word_generator is None:
        try:
            word_generator = WordGenerator()
        except Exception as e:
            logger.error(f"Error creando WordGenerator: {e}")
    return word_generator