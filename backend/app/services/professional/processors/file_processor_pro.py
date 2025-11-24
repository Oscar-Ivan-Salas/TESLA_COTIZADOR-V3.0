"""
PROCESADOR DE ARCHIVOS PROFESIONAL v4.0
Extraccion de texto de multiples formatos con OCR

Formatos soportados:
- PDF (texto y escaneados)
- Word (.docx, .doc)
- Excel (.xlsx, .xls, .csv)
- Imagenes (PNG, JPG, TIFF) con OCR
- Texto plano (TXT, JSON, XML)
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import tempfile
import base64

logger = logging.getLogger(__name__)

# Imports condicionales para manejo de errores
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber no disponible - pip install pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no disponible - pip install python-docx")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas no disponible - pip install pandas")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("OCR no disponible - pip install pillow pytesseract")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl no disponible - pip install openpyxl")


class FileProcessorPro:
    """
    Procesador profesional de archivos con soporte multi-formato.

    Extrae texto, tablas y metadatos de diversos tipos de archivos
    para alimentar el sistema RAG y generacion de documentos.
    """

    def __init__(self, upload_dir: str = None):
        """
        Inicializa el procesador de archivos.

        Args:
            upload_dir: Directorio para archivos subidos
        """
        self.upload_dir = Path(upload_dir) if upload_dir else Path("backend/storage/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        # Estadisticas de capacidades
        self.capabilities = {
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "excel": PANDAS_AVAILABLE and EXCEL_AVAILABLE,
            "ocr": OCR_AVAILABLE,
            "csv": PANDAS_AVAILABLE,
            "json": True,
            "txt": True
        }

        logger.info(f"FileProcessorPro inicializado - Capacidades: {self.capabilities}")

    def process_file(
        self,
        file_path: Union[str, Path],
        extract_tables: bool = True,
        ocr_enabled: bool = True
    ) -> Dict[str, Any]:
        """
        Procesa un archivo y extrae su contenido.

        Args:
            file_path: Ruta al archivo
            extract_tables: Extraer tablas si es posible
            ocr_enabled: Usar OCR para imagenes/PDFs escaneados

        Returns:
            Dict con texto, tablas, metadatos
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                "success": False,
                "error": f"Archivo no encontrado: {file_path}",
                "text": "",
                "tables": [],
                "metadata": {}
            }

        extension = file_path.suffix.lower()

        try:
            # Seleccionar procesador segun extension
            if extension == '.pdf':
                return self._process_pdf(file_path, extract_tables, ocr_enabled)
            elif extension in ['.docx', '.doc']:
                return self._process_word(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._process_excel(file_path)
            elif extension == '.csv':
                return self._process_csv(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._process_image(file_path, ocr_enabled)
            elif extension == '.json':
                return self._process_json(file_path)
            elif extension in ['.txt', '.md', '.rst']:
                return self._process_text(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Formato no soportado: {extension}",
                    "text": "",
                    "tables": [],
                    "metadata": {"extension": extension}
                }

        except Exception as e:
            logger.error(f"Error procesando {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "tables": [],
                "metadata": {"file": str(file_path)}
            }

    def process_multiple(
        self,
        file_paths: List[Union[str, Path]],
        extract_tables: bool = True,
        ocr_enabled: bool = True
    ) -> Dict[str, Any]:
        """
        Procesa multiples archivos y combina resultados.

        Returns:
            Dict con texto combinado y lista de resultados individuales
        """
        results = []
        combined_text = []
        all_tables = []

        for file_path in file_paths:
            result = self.process_file(file_path, extract_tables, ocr_enabled)
            results.append(result)

            if result.get("success", False):
                combined_text.append(result.get("text", ""))
                all_tables.extend(result.get("tables", []))

        return {
            "success": True,
            "total_files": len(file_paths),
            "processed": sum(1 for r in results if r.get("success")),
            "combined_text": "\n\n---\n\n".join(combined_text),
            "all_tables": all_tables,
            "individual_results": results
        }

    def process_base64(
        self,
        base64_content: str,
        filename: str,
        extract_tables: bool = True,
        ocr_enabled: bool = True
    ) -> Dict[str, Any]:
        """
        Procesa un archivo desde base64.

        Util para archivos subidos via API.
        """
        try:
            # Decodificar base64
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]

            file_bytes = base64.b64decode(base64_content)

            # Guardar temporalmente
            temp_path = self.upload_dir / f"temp_{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"

            with open(temp_path, 'wb') as f:
                f.write(file_bytes)

            # Procesar
            result = self.process_file(temp_path, extract_tables, ocr_enabled)

            # Limpiar archivo temporal
            temp_path.unlink()

            return result

        except Exception as e:
            logger.error(f"Error procesando base64: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "tables": [],
                "metadata": {}
            }

    # =========================================================================
    # PROCESADORES ESPECIFICOS
    # =========================================================================

    def _process_pdf(
        self,
        file_path: Path,
        extract_tables: bool = True,
        ocr_enabled: bool = True
    ) -> Dict[str, Any]:
        """Procesa archivo PDF con extraccion de texto y tablas"""

        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "pdfplumber no instalado",
                "text": "",
                "tables": [],
                "metadata": {}
            }

        text_content = []
        tables = []

        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "pages": len(pdf.pages),
                "filename": file_path.name,
                "type": "pdf"
            }

            for i, page in enumerate(pdf.pages):
                # Extraer texto
                page_text = page.extract_text()

                if page_text:
                    text_content.append(f"--- Pagina {i+1} ---\n{page_text}")
                elif ocr_enabled and OCR_AVAILABLE:
                    # Intentar OCR si no hay texto
                    img = page.to_image(resolution=300)
                    ocr_text = pytesseract.image_to_string(img.original, lang='spa')
                    if ocr_text.strip():
                        text_content.append(f"--- Pagina {i+1} (OCR) ---\n{ocr_text}")

                # Extraer tablas
                if extract_tables:
                    page_tables = page.extract_tables()
                    for j, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                "page": i + 1,
                                "table_index": j,
                                "data": table
                            })

        return {
            "success": True,
            "text": "\n\n".join(text_content),
            "tables": tables,
            "metadata": metadata
        }

    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo Word (.docx)"""

        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx no instalado",
                "text": "",
                "tables": [],
                "metadata": {}
            }

        doc = Document(file_path)

        # Extraer texto de parrafos
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # Extraer tablas
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                tables.append({
                    "table_index": i,
                    "data": table_data
                })

        return {
            "success": True,
            "text": "\n\n".join(text_content),
            "tables": tables,
            "metadata": {
                "filename": file_path.name,
                "type": "word",
                "paragraphs": len(text_content),
                "tables": len(tables)
            }
        }

    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo Excel (.xlsx, .xls)"""

        if not PANDAS_AVAILABLE:
            return {
                "success": False,
                "error": "pandas no instalado",
                "text": "",
                "tables": [],
                "metadata": {}
            }

        # Leer todas las hojas
        excel_file = pd.ExcelFile(file_path)

        text_content = []
        tables = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            # Convertir a texto
            sheet_text = f"=== Hoja: {sheet_name} ===\n"
            sheet_text += df.to_string()
            text_content.append(sheet_text)

            # Guardar como tabla
            tables.append({
                "sheet": sheet_name,
                "data": df.values.tolist(),
                "columns": df.columns.tolist(),
                "rows": len(df)
            })

        return {
            "success": True,
            "text": "\n\n".join(text_content),
            "tables": tables,
            "metadata": {
                "filename": file_path.name,
                "type": "excel",
                "sheets": len(excel_file.sheet_names)
            }
        }

    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo CSV"""

        if not PANDAS_AVAILABLE:
            # Fallback sin pandas
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            return {
                "success": True,
                "text": content,
                "tables": [],
                "metadata": {"filename": file_path.name, "type": "csv"}
            }

        # Con pandas
        df = pd.read_csv(file_path)

        return {
            "success": True,
            "text": df.to_string(),
            "tables": [{
                "data": df.values.tolist(),
                "columns": df.columns.tolist(),
                "rows": len(df)
            }],
            "metadata": {
                "filename": file_path.name,
                "type": "csv",
                "columns": len(df.columns),
                "rows": len(df)
            }
        }

    def _process_image(self, file_path: Path, ocr_enabled: bool = True) -> Dict[str, Any]:
        """Procesa imagen con OCR"""

        if not ocr_enabled:
            return {
                "success": True,
                "text": "[Imagen sin procesar - OCR deshabilitado]",
                "tables": [],
                "metadata": {"filename": file_path.name, "type": "image"}
            }

        if not OCR_AVAILABLE:
            return {
                "success": False,
                "error": "OCR no disponible - instalar pytesseract y PIL",
                "text": "",
                "tables": [],
                "metadata": {}
            }

        # Abrir imagen
        image = Image.open(file_path)

        # Realizar OCR
        text = pytesseract.image_to_string(image, lang='spa')

        return {
            "success": True,
            "text": text,
            "tables": [],
            "metadata": {
                "filename": file_path.name,
                "type": "image",
                "size": image.size,
                "mode": image.mode
            }
        }

    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo JSON"""

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convertir a texto legible
        text = json.dumps(data, indent=2, ensure_ascii=False)

        return {
            "success": True,
            "text": text,
            "tables": [],
            "metadata": {
                "filename": file_path.name,
                "type": "json",
                "keys": list(data.keys()) if isinstance(data, dict) else []
            },
            "json_data": data
        }

    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo de texto plano"""

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        return {
            "success": True,
            "text": text,
            "tables": [],
            "metadata": {
                "filename": file_path.name,
                "type": "text",
                "lines": len(text.split('\n')),
                "characters": len(text)
            }
        }

    # =========================================================================
    # METODOS AUXILIARES
    # =========================================================================

    def get_capabilities(self) -> Dict[str, bool]:
        """Retorna las capacidades disponibles del procesador"""
        return self.capabilities

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[str]:
        """
        Divide texto en chunks para el sistema RAG.

        Args:
            text: Texto a dividir
            chunk_size: Tamano aproximado de cada chunk (en palabras)
            overlap: Palabras de superposicion entre chunks

        Returns:
            Lista de chunks de texto
        """
        words = text.split()
        chunks = []

        if len(words) <= chunk_size:
            return [text]

        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk = ' '.join(words[start:end])
            chunks.append(chunk)
            start = end - overlap

        return chunks


# Instancia global
file_processor_pro = FileProcessorPro()

def get_file_processor() -> FileProcessorPro:
    """Obtiene la instancia del procesador de archivos"""
    return file_processor_pro
