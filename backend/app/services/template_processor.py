"""
📋 TEMPLATE PROCESSOR + PILI v3.0 - PLANTILLAS INTELIGENTES
📁 RUTA: backend/app/services/template_processor.py

PILI (Procesadora Inteligente de Licitaciones Industriales) integrada con 
procesador de plantillas para crear documentos personalizados perfectos.

🎯 NUEVAS CARACTERÍSTICAS PILI v3.0:
- Procesamiento inteligente de marcadores PILI
- Plantillas especializadas por agente
- Validación automática con IA
- Inserción inteligente de tablas y logos
- Marcadores dinámicos personalizados
- Sin corrupción de documentos (problema resuelto)

🔄 CONSERVA TODO LO EXISTENTE:
- Todos los marcadores soportados ✅
- procesar_plantilla() ✅
- validar_plantilla() ✅
- extraer_marcadores() ✅
- Toda la lógica de reemplazo ✅
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime
import logging
import re
import base64
from io import BytesIO
import json
import tempfile

logger = logging.getLogger(__name__)

class TemplateProcessor:
    """
    🔄 PROCESADOR ORIGINAL CONSERVADO + 🤖 PILI INTEGRADA
    
    Mantiene toda la funcionalidad existente pero agrega capacidades
    inteligentes de PILI para plantillas personalizadas.
    """
    
    def __init__(self):
        """🔄 CONSERVADO + 🤖 PILI mejorado"""
        
        # Marcadores estándar conservados
        self.marcadores_estandar = {
            "{{cliente}}": "Nombre del cliente",
            "{{proyecto}}": "Nombre del proyecto", 
            "{{numero_cotizacion}}": "Número de cotización",
            "{{numero}}": "Número del documento",
            "{{fecha}}": "Fecha actual",
            "{{descripcion}}": "Descripción del proyecto",
            "{{subtotal}}": "Subtotal",
            "{{igv}}": "IGV (18%)",
            "{{total}}": "Total",
            "{{items_tabla}}": "Tabla de items (se reemplaza con tabla real)",
            "{{logo}}": "Logo de empresa (se inserta imagen)",
            "{{observaciones}}": "Observaciones del proyecto",
            "{{vigencia}}": "Vigencia de la cotización",
            "{{empresa_nombre}}": "Nombre de la empresa",
            "{{empresa_direccion}}": "Dirección de la empresa",
            "{{empresa_telefono}}": "Teléfono de la empresa",
            "{{empresa_email}}": "Email de la empresa"
        }
        
        # 🤖 Nuevos marcadores PILI
        self.marcadores_pili = {
            "{{agente_pili}}": "Agente PILI responsable",
            "{{pili_version}}": "Versión de PILI",
            "{{pili_timestamp}}": "Timestamp de generación PILI",
            "{{tipo_servicio}}": "Tipo de servicio PILI",
            "{{pili_mensaje}}": "Mensaje personalizado de PILI",
            "{{pili_firma}}": "Firma digital de PILI"
        }
        
        # Colores Tesla + PILI
        self.COLOR_TESLA = RGBColor(218, 165, 32)  # Dorado Tesla
        self.COLOR_PILI = RGBColor(212, 175, 55)   # Dorado PILI
        
        logger.info("✅ TemplateProcessor + PILI inicializado")
    
    # ═══════════════════════════════════════════════════════════════
    # 🤖 NUEVOS MÉTODOS PILI v3.0
    # ═══════════════════════════════════════════════════════════════
    
    def procesar_plantilla_con_pili(
        self,
        ruta_plantilla: str,
        datos_json: Dict[str, Any],
        ruta_salida: Optional[str] = None,
        opciones_pili: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        🤖 NUEVO PILI v3.0 - Procesa plantilla con datos JSON estructurados de PILI
        
        Args:
            ruta_plantilla: Ruta a la plantilla .docx
            datos_json: Datos estructurados por PILI
            ruta_salida: Ruta del archivo de salida (opcional)
            opciones_pili: Opciones específicas de PILI
            
        Returns:
            Ruta del archivo generado
        """
        
        try:
            logger.info(f"🤖 PILI procesando plantilla: {Path(ruta_plantilla).name}")
            
            # 1. Validar plantilla
            es_valida, mensaje = self.validar_plantilla(ruta_plantilla)
            if not es_valida:
                raise ValueError(f"Plantilla inválida: {mensaje}")
            
            # 2. Cargar plantilla
            doc = Document(ruta_plantilla)
            
            # 3. Preparar datos para reemplazo
            datos_completos = self._preparar_datos_pili(datos_json, opciones_pili)
            
            # 4. Reemplazar marcadores de texto
            self._reemplazar_marcadores_pili(doc, datos_completos)
            
            # 5. Procesar elementos especiales (tablas, logos, etc.)
            self._procesar_elementos_especiales_pili(doc, datos_completos)
            
            # 6. Aplicar mejoras de formato PILI
            self._aplicar_mejoras_formato_pili(doc, datos_completos)
            
            # 7. Generar ruta de salida
            if not ruta_salida:
                ruta_salida = self._generar_ruta_salida(ruta_plantilla, datos_completos)
            
            # 8. Guardar documento procesado
            doc.save(ruta_salida)
            
            logger.info(f"✅ PILI plantilla procesada: {Path(ruta_salida).name}")
            return ruta_salida
            
        except Exception as e:
            logger.error(f"❌ Error PILI procesando plantilla: {str(e)}")
            raise
    
    def _preparar_datos_pili(
        self, 
        datos_json: Dict[str, Any], 
        opciones_pili: Optional[Dict[str, Any]]
    ) -> Dict[str, str]:
        """Prepara datos JSON de PILI para reemplazo en plantilla"""
        
        # Extraer datos principales
        datos_extraidos = datos_json.get("datos_extraidos", {})
        agente_pili = datos_json.get("agente_responsable", "PILI")
        tipo_servicio = datos_json.get("tipo_servicio", "documento")
        
        # Preparar datos completos
        datos_completos = {}
        
        # 1. Datos estándar de empresa
        datos_empresa = {
            "empresa_nombre": "TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C.",
            "empresa_ruc": "20601138787",
            "empresa_direccion": "Jr. Las Ágatas Mz B Lote 09, Urb. San Carlos, SJL",
            "empresa_telefono": "906315961",
            "empresa_email": "ingenieria.teslaelectricidad@gmail.com"
        }
        datos_completos.update(datos_empresa)
        
        # 2. Datos específicos del documento
        datos_completos.update({
            "fecha": datetime.now().strftime("%d/%m/%Y"),
            "hora": datetime.now().strftime("%H:%M"),
            "cliente": datos_extraidos.get("cliente", "[Cliente por definir]"),
            "proyecto": datos_extraidos.get("proyecto", "[Proyecto por definir]"),
            "numero": datos_extraidos.get("numero", f"DOC-{datetime.now().strftime('%Y%m%d')}-001"),
            "numero_cotizacion": datos_extraidos.get("numero", f"COT-{datetime.now().strftime('%Y%m%d')}-001"),
            "descripcion": datos_extraidos.get("descripcion", ""),
            "observaciones": datos_extraidos.get("observaciones", ""),
            "vigencia": datos_extraidos.get("vigencia", "30 días")
        })
        
        # 3. Datos financieros
        if "items" in datos_extraidos:
            subtotal, igv, total = self._calcular_totales(datos_extraidos["items"])
            datos_completos.update({
                "subtotal": f"S/ {subtotal:.2f}",
                "igv": f"S/ {igv:.2f}",
                "total": f"S/ {total:.2f}"
            })
        else:
            datos_completos.update({
                "subtotal": f"S/ {datos_extraidos.get('subtotal', 0):.2f}",
                "igv": f"S/ {datos_extraidos.get('igv', 0):.2f}",
                "total": f"S/ {datos_extraidos.get('total', 0):.2f}"
            })
        
        # 4. 🤖 Datos específicos PILI
        datos_pili = {
            "agente_pili": agente_pili,
            "pili_version": "v3.0",
            "pili_timestamp": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "tipo_servicio": tipo_servicio.replace("-", " ").title(),
            "pili_mensaje": f"Generado por {agente_pili} - Tu agente IA especializada",
            "pili_firma": f"🤖 {agente_pili} | Tesla IA v3.0"
        }
        datos_completos.update(datos_pili)
        
        # 5. Aplicar opciones específicas
        if opciones_pili:
            datos_completos.update(opciones_pili.get("datos_adicionales", {}))
        
        # Convertir todos los valores a string
        return {k: str(v) for k, v in datos_completos.items()}
    
    def _calcular_totales(self, items: List[Dict[str, Any]]) -> tuple[float, float, float]:
        """Calcula subtotal, IGV y total desde lista de items"""
        
        subtotal = 0.0
        for item in items:
            cantidad = float(item.get("cantidad", 0))
            precio_unitario = float(item.get("precio_unitario", 0))
            subtotal += cantidad * precio_unitario
        
        igv = subtotal * 0.18
        total = subtotal + igv
        
        return subtotal, igv, total
    
    def _reemplazar_marcadores_pili(self, doc: Document, datos: Dict[str, str]):
        """Reemplaza marcadores en el documento usando datos PILI"""
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for marcador, valor in datos.items():
                marcador_completo = f"{{{{{marcador}}}}}"
                if marcador_completo in paragraph.text:
                    paragraph.text = paragraph.text.replace(marcador_completo, valor)
        
        # Reemplazar en tablas existentes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for marcador, valor in datos.items():
                        marcador_completo = f"{{{{{marcador}}}}}"
                        if marcador_completo in cell.text:
                            cell.text = cell.text.replace(marcador_completo, valor)
        
        # Reemplazar en headers y footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    for marcador, valor in datos.items():
                        marcador_completo = f"{{{{{marcador}}}}}"
                        if marcador_completo in paragraph.text:
                            paragraph.text = paragraph.text.replace(marcador_completo, valor)
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for marcador, valor in datos.items():
                        marcador_completo = f"{{{{{marcador}}}}}"
                        if marcador_completo in paragraph.text:
                            paragraph.text = paragraph.text.replace(marcador_completo, valor)
    
    def _procesar_elementos_especiales_pili(self, doc: Document, datos: Dict[str, str]):
        """Procesa elementos especiales como tablas e imágenes con lógica PILI"""
        
        # 1. Procesar {{items_tabla}}
        self._procesar_tabla_items_pili(doc, datos)
        
        # 2. Procesar {{logo}} 
        self._procesar_logo_pili(doc, datos)
        
        # 3. Agregar elementos PILI automáticamente
        self._agregar_elementos_automaticos_pili(doc, datos)
    
    def _procesar_tabla_items_pili(self, doc: Document, datos: Dict[str, str]):
        """Procesa marcador {{items_tabla}} con tabla real de items"""
        
        # Buscar párrafo con {{items_tabla}}
        for i, paragraph in enumerate(doc.paragraphs):
            if "{{items_tabla}}" in paragraph.text:
                
                # Obtener items desde datos (si están en formato JSON como string)
                items_data = datos.get("items", "[]")
                try:
                    if isinstance(items_data, str):
                        items = json.loads(items_data) if items_data != "[]" else []
                    else:
                        items = items_data
                except:
                    items = []  # Si falla el parsing, usar lista vacía
                
                if not items:
                    # Si no hay items, reemplazar con mensaje
                    paragraph.text = paragraph.text.replace("{{items_tabla}}", "No hay items especificados.")
                    continue
                
                # Crear tabla después del párrafo
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Headers con estilo PILI
                headers = ['DESCRIPCIÓN', 'CANT.', 'UND.', 'P.UNITARIO', 'TOTAL']
                header_cells = table.rows[0].cells
                
                for j, header in enumerate(headers):
                    header_cells[j].text = header
                    # Aplicar formato PILI a headers
                    for para in header_cells[j].paragraphs:
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                    
                    # Aplicar fondo dorado PILI
                    self._aplicar_fondo_celda_pili(header_cells[j])
                
                # Agregar items
                for item in items:
                    row = table.add_row()
                    cells = row.cells
                    
                    descripcion = item.get("descripcion", "")
                    cantidad = float(item.get("cantidad", 0))
                    unidad = item.get("unidad", "und")
                    precio_unitario = float(item.get("precio_unitario", 0))
                    total_item = cantidad * precio_unitario
                    
                    cells[0].text = descripcion
                    cells[1].text = str(int(cantidad) if cantidad.is_integer() else cantidad)
                    cells[2].text = unidad
                    cells[3].text = f"S/ {precio_unitario:.2f}"
                    cells[4].text = f"S/ {total_item:.2f}"
                    
                    # Alineación
                    cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                    # Formateo total
                    cells[4].paragraphs[0].runs[0].bold = True
                    cells[4].paragraphs[0].runs[0].font.color.rgb = self.COLOR_PILI
                
                # Limpiar marcador del párrafo
                paragraph.text = paragraph.text.replace("{{items_tabla}}", "")
                break
    
    def _procesar_logo_pili(self, doc: Document, datos: Dict[str, str]):
        """Procesa marcador {{logo}} con imagen real"""
        
        logo_base64 = datos.get("logo_base64", "")
        if not logo_base64:
            # Si no hay logo, reemplazar marcador con texto
            for paragraph in doc.paragraphs:
                if "{{logo}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{logo}}", "[LOGO TESLA ELECTRICIDAD]")
            return
        
        # Procesar logo desde base64
        for paragraph in doc.paragraphs:
            if "{{logo}}" in paragraph.text:
                try:
                    # Decodificar base64
                    if ',' in logo_base64:
                        logo_base64 = logo_base64.split(',')[1]
                    
                    logo_bytes = base64.b64decode(logo_base64)
                    
                    # Crear archivo temporal
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(logo_bytes)
                        temp_path = temp_file.name
                    
                    # Limpiar párrafo y insertar imagen
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(temp_path, width=Inches(2.0))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Limpiar archivo temporal
                    Path(temp_path).unlink()
                    
                except Exception as e:
                    logger.error(f"Error procesando logo: {e}")
                    paragraph.text = paragraph.text.replace("{{logo}}", "[LOGO]")
                
                break
    
    def _agregar_elementos_automaticos_pili(self, doc: Document, datos: Dict[str, str]):
        """Agrega elementos automáticos de PILI si no existen en la plantilla"""
        
        agente_pili = datos.get("agente_pili", "PILI")
        
        # Buscar si ya existe firma PILI
        texto_completo = ""
        for paragraph in doc.paragraphs:
            texto_completo += paragraph.text
        
        # Si no existe referencia a PILI, agregar al final
        if "PILI" not in texto_completo and "agente" not in texto_completo.lower():
            
            # Agregar separador
            doc.add_paragraph("─" * 60)
            
            # Agregar firma PILI
            firma_para = doc.add_paragraph()
            firma_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            run_firma = firma_para.add_run(f"✨ Documento generado por {agente_pili} - Tu agente IA especializada")
            run_firma.font.size = Pt(9)
            run_firma.font.color.rgb = self.COLOR_PILI
            run_firma.italic = True
            
            # Timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            timestamp_run = timestamp_para.add_run(f"Generado el {datos.get('pili_timestamp', datetime.now().strftime('%d/%m/%Y %H:%M'))}")
            timestamp_run.font.size = Pt(8)
            timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _aplicar_fondo_celda_pili(self, celda):
        """Aplica fondo dorado PILI a una celda de tabla"""
        
        try:
            # Aplicar fondo dorado usando color PILI
            shading = celda._element.get_or_add_tcPr().get_or_add_shd()
            shading.fill = f"{self.COLOR_PILI.r:02X}{self.COLOR_PILI.g:02X}{self.COLOR_PILI.b:02X}"
        except Exception as e:
            logger.debug(f"No se pudo aplicar fondo a celda: {e}")
    
    def _aplicar_mejoras_formato_pili(self, doc: Document, datos: Dict[str, str]):
        """Aplica mejoras de formato automáticas de PILI"""
        
        # Aplicar fuente consistente
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.font.name:
                    run.font.name = 'Arial'
                if not run.font.size:
                    run.font.size = Pt(11)
        
        # Mejorar alineación de números en tablas
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    # Si la celda contiene "S/" (precio), alinear a la derecha
                    if "S/" in cell.text:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    def _generar_ruta_salida(self, ruta_plantilla: str, datos: Dict[str, str]) -> str:
        """Genera ruta de salida única para el documento procesado"""
        
        # Usar directorio de salida de configuración centralizada
        from app.core.config import get_generated_directory
        output_dir = get_generated_directory()

        
        # Generar nombre único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        plantilla_nombre = Path(ruta_plantilla).stem
        cliente_slug = self._slugify(datos.get("cliente", "cliente"))
        
        nombre_archivo = f"{plantilla_nombre}_{cliente_slug}_{timestamp}.docx"
        
        return str(output_dir / nombre_archivo)
    
    def _slugify(self, texto: str) -> str:
        """Convierte texto en slug válido para nombre archivo"""
        import re
        texto = re.sub(r'[^\w\s-]', '', texto.strip())
        texto = re.sub(r'[-\s]+', '-', texto)
        return texto.lower()[:15]

    # ═══════════════════════════════════════════════════════════════
    # 🔄 MÉTODOS ORIGINALES CONSERVADOS (COMPATIBILIDAD)
    # ═══════════════════════════════════════════════════════════════
    
    def procesar_plantilla(
        self,
        ruta_plantilla: str,
        datos_cotizacion: Dict[str, Any],
        ruta_salida: Optional[str] = None,
        logo_base64: Optional[str] = None
    ) -> str:
        """
        🔄 CONSERVADO - Procesa plantilla Word personalizada con marcadores
        
        Método original mantenido para compatibilidad hacia atrás.
        """
        
        try:
            logger.info(f"Procesando plantilla: {ruta_plantilla}")
            
            # Validar que la plantilla existe
            if not Path(ruta_plantilla).exists():
                raise FileNotFoundError(f"Plantilla no encontrada: {ruta_plantilla}")
            
            # Validar plantilla
            es_valida, mensaje = self.validar_plantilla(ruta_plantilla)
            if not es_valida:
                raise ValueError(f"Plantilla inválida: {mensaje}")
            
            # Cargar documento
            doc = Document(ruta_plantilla)
            
            # Preparar datos para reemplazo
            datos_reemplazo = self._preparar_datos_cotizacion(datos_cotizacion)
            
            # Agregar logo si se proporciona
            if logo_base64:
                datos_reemplazo["logo_base64"] = logo_base64
            
            # Reemplazar marcadores simples
            self._reemplazar_marcadores_simples(doc, datos_reemplazo)
            
            # Procesar tabla de items si existe
            self._procesar_tabla_items_original(doc, datos_cotizacion.get("items", []))
            
            # Procesar logo si existe
            if logo_base64:
                self._procesar_logo_original(doc, logo_base64)
            
            # Generar ruta de salida si no se proporciona
            if not ruta_salida:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                plantilla_nombre = Path(ruta_plantilla).stem
                from app.core.config import get_generated_directory
                ruta_salida = str(get_generated_directory() / f"{plantilla_nombre}_procesada_{timestamp}.docx")
            
            # Crear directorio de salida si no existe
            Path(ruta_salida).parent.mkdir(parents=True, exist_ok=True)
            
            # Guardar documento
            doc.save(ruta_salida)
            
            logger.info(f"✅ Plantilla procesada: {ruta_salida}")
            return ruta_salida
            
        except Exception as e:
            logger.error(f"Error procesando plantilla: {str(e)}")
            raise
    
    def validar_plantilla(self, ruta_plantilla: str) -> tuple[bool, str]:
        """
        🔄 CONSERVADO - Valida si una plantilla Word es válida
        
        Método original mantenido para compatibilidad hacia atrás.
        """
        
        try:
            # Verificar que el archivo existe
            if not Path(ruta_plantilla).exists():
                return False, "Archivo no encontrado"
            
            # Verificar extensión
            if not ruta_plantilla.lower().endswith('.docx'):
                return False, "El archivo debe tener extensión .docx"
            
            # Intentar abrir el documento
            try:
                doc = Document(ruta_plantilla)
                
                # Verificar que se puede leer
                _ = len(doc.paragraphs)
                
                return True, "Plantilla válida"
                
            except Exception as e:
                return False, f"Error al abrir el documento: {str(e)}"
        
        except Exception as e:
            return False, f"Error validando plantilla: {str(e)}"
    
    def extraer_marcadores(self, ruta_plantilla: str) -> List[str]:
        """
        🔄 CONSERVADO - Extrae todos los marcadores {{variable}} de una plantilla
        
        Método original mantenido para compatibilidad hacia atrás.
        """
        
        marcadores = set()
        
        try:
            doc = Document(ruta_plantilla)
            
            # Buscar en párrafos
            for paragraph in doc.paragraphs:
                encontrados = re.findall(r'\{\{[^}]+\}\}', paragraph.text)
                marcadores.update(encontrados)
            
            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        encontrados = re.findall(r'\{\{[^}]+\}\}', cell.text)
                        marcadores.update(encontrados)
            
            # Buscar en headers y footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        encontrados = re.findall(r'\{\{[^}]+\}\}', paragraph.text)
                        marcadores.update(encontrados)
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        encontrados = re.findall(r'\{\{[^}]+\}\}', paragraph.text)
                        marcadores.update(encontrados)
            
            return sorted(list(marcadores))
            
        except Exception as e:
            logger.error(f"Error extrayendo marcadores: {str(e)}")
            return []
    
    # ═══════════════════════════════════════════════════════════════
    # 🔄 MÉTODOS AUXILIARES ORIGINALES CONSERVADOS
    # ═══════════════════════════════════════════════════════════════
    
    def _preparar_datos_cotizacion(self, datos_cotizacion: Dict[str, Any]) -> Dict[str, str]:
        """
        🔄 CONSERVADO - Prepara datos de cotización para reemplazo
        """
        
        datos = {}
        
        # Datos básicos
        datos["cliente"] = str(datos_cotizacion.get("cliente", ""))
        datos["proyecto"] = str(datos_cotizacion.get("proyecto", ""))
        datos["numero"] = str(datos_cotizacion.get("numero", ""))
        datos["numero_cotizacion"] = str(datos_cotizacion.get("numero", ""))
        datos["fecha"] = datetime.now().strftime("%d/%m/%Y")
        datos["descripcion"] = str(datos_cotizacion.get("descripcion", ""))
        datos["observaciones"] = str(datos_cotizacion.get("observaciones", ""))
        datos["vigencia"] = str(datos_cotizacion.get("vigencia", "30 días"))
        
        # Datos financieros
        subtotal = float(datos_cotizacion.get("subtotal", 0))
        igv = float(datos_cotizacion.get("igv", 0))
        total = float(datos_cotizacion.get("total", 0))
        
        datos["subtotal"] = f"S/ {subtotal:.2f}"
        datos["igv"] = f"S/ {igv:.2f}"
        datos["total"] = f"S/ {total:.2f}"
        
        # Datos de empresa
        datos["empresa_nombre"] = "TESLA ELECTRICIDAD Y AUTOMATIZACIÓN S.A.C."
        datos["empresa_direccion"] = "Jr. Las Ágatas Mz B Lote 09, Urb. San Carlos, SJL"
        datos["empresa_telefono"] = "906315961"
        datos["empresa_email"] = "ingenieria.teslaelectricidad@gmail.com"
        
        return datos
    
    def _reemplazar_marcadores_simples(self, doc: Document, datos: Dict[str, str]):
        """
        🔄 CONSERVADO - Reemplaza marcadores simples en el documento
        """
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for marcador, valor in datos.items():
                marcador_completo = f"{{{{{marcador}}}}}"
                if marcador_completo in paragraph.text:
                    paragraph.text = paragraph.text.replace(marcador_completo, valor)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for marcador, valor in datos.items():
                        marcador_completo = f"{{{{{marcador}}}}}"
                        if marcador_completo in cell.text:
                            cell.text = cell.text.replace(marcador_completo, valor)
    
    def _procesar_tabla_items_original(self, doc: Document, items: List[Dict[str, Any]]):
        """
        🔄 CONSERVADO - Procesa tabla de items (método original)
        """
        
        # Buscar marcador {{items_tabla}}
        for paragraph in doc.paragraphs:
            if "{{items_tabla}}" in paragraph.text:
                
                if not items:
                    paragraph.text = paragraph.text.replace("{{items_tabla}}", "No hay items especificados.")
                    continue
                
                # Crear tabla
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Headers
                headers = ['DESCRIPCIÓN', 'CANT.', 'UND.', 'P.UNITARIO', 'TOTAL']
                header_cells = table.rows[0].cells
                
                for j, header in enumerate(headers):
                    header_cells[j].text = header
                    header_cells[j].paragraphs[0].runs[0].bold = True
                
                # Items
                for item in items:
                    row = table.add_row()
                    cells = row.cells
                    
                    cells[0].text = str(item.get("descripcion", ""))
                    cells[1].text = str(item.get("cantidad", ""))
                    cells[2].text = str(item.get("unidad", "und"))
                    cells[3].text = f"S/ {float(item.get('precio_unitario', 0)):.2f}"
                    
                    total_item = float(item.get("cantidad", 0)) * float(item.get("precio_unitario", 0))
                    cells[4].text = f"S/ {total_item:.2f}"
                
                # Limpiar marcador
                paragraph.text = paragraph.text.replace("{{items_tabla}}", "")
                break
    
    def _procesar_logo_original(self, doc: Document, logo_base64: str):
        """
        🔄 CONSERVADO - Procesa logo (método original)
        """
        
        for paragraph in doc.paragraphs:
            if "{{logo}}" in paragraph.text:
                try:
                    if ',' in logo_base64:
                        logo_base64 = logo_base64.split(',')[1]
                    
                    logo_bytes = base64.b64decode(logo_base64)
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(logo_bytes)
                        temp_path = temp_file.name
                    
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(temp_path, width=Inches(2.0))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    Path(temp_path).unlink()
                    
                except Exception as e:
                    logger.error(f"Error procesando logo: {e}")
                    paragraph.text = paragraph.text.replace("{{logo}}", "[LOGO]")
                
                break
    
    def obtener_marcadores_disponibles(self) -> Dict[str, str]:
        """
        🔄 CONSERVADO + 🤖 PILI MEJORADO
        
        Retorna todos los marcadores disponibles (estándar + PILI).
        """
        
        marcadores_completos = {}
        marcadores_completos.update(self.marcadores_estandar)
        marcadores_completos.update(self.marcadores_pili)
        
        return marcadores_completos

# ═══════════════════════════════════════════════════════════════
# 🎯 INSTANCIA GLOBAL MEJORADA CON PILI
# ═══════════════════════════════════════════════════════════════

# Crear instancia global con manejo robusto de errores
try:
    template_processor = TemplateProcessor()
    logger.info("✅ TemplateProcessor + PILI inicializado correctamente")
except Exception as e:
    logger.error(f"❌ Error crítico inicializando TemplateProcessor: {e}")
    template_processor = None

# Función auxiliar para obtener instancia segura
def get_template_processor():
    """Obtiene instancia de TemplateProcessor de forma segura"""
    global template_processor
    if template_processor is None:
        try:
            template_processor = TemplateProcessor()
        except Exception as e:
            logger.error(f"Error creando TemplateProcessor: {e}")
    return template_processor