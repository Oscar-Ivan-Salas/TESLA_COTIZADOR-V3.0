#!/usr/bin/env python3
"""
🧪 SCRIPT DE PRUEBA: Generación de Documentos Word
Verifica que el sistema genera correctamente documentos Word para cotizaciones y proyectos
"""

import sys
import os
from pathlib import Path
from datetime import datetime

# Agregar el directorio del proyecto al path
sys.path.insert(0, str(Path(__file__).parent))

from app.services.word_generator import word_generator
from app.core.config import settings

def print_section(title):
    """Imprimir sección con formato"""
    print("\n" + "="*70)
    print(f"  {title}")
    print("="*70)

def test_generar_cotizacion_word():
    """Test 1: Generar cotización en Word"""
    print_section("TEST 1: Generación de Cotización Word")

    try:
        # Datos de prueba para cotización
        datos_cotizacion = {
            "numero": "COT-202512-TEST-001",
            "fecha": datetime.now().strftime("%d/%m/%Y"),
            "cliente": "CLIENTE DE PRUEBA S.A.C.",
            "proyecto": "Instalación Eléctrica Oficinas - PRUEBA",
            "descripcion": "Cotización de prueba generada automáticamente",
            "vigencia": "30 días",
            "items": [
                {
                    "descripcion": "Instalación de puntos de luz LED 18W",
                    "cantidad": 20,
                    "unidad": "und",
                    "precio_unitario": 85.00,
                    "total": 1700.00
                },
                {
                    "descripcion": "Instalación de tomacorrientes dobles empotrados",
                    "cantidad": 15,
                    "unidad": "und",
                    "precio_unitario": 65.00,
                    "total": 975.00
                },
                {
                    "descripcion": "Tablero eléctrico trifásico 24 polos",
                    "cantidad": 1,
                    "unidad": "und",
                    "precio_unitario": 850.00,
                    "total": 850.00
                },
                {
                    "descripcion": "Cable NYY 3x6mm² (por metro)",
                    "cantidad": 50,
                    "unidad": "m",
                    "precio_unitario": 12.50,
                    "total": 625.00
                },
                {
                    "descripcion": "Tubo PVC SEL 25mm (por metro)",
                    "cantidad": 50,
                    "unidad": "m",
                    "precio_unitario": 4.50,
                    "total": 225.00
                }
            ],
            "subtotal": 4375.00,
            "igv": 787.50,
            "total": 5162.50,
            "observaciones": "Esta es una cotización de prueba para verificar la generación correcta de documentos Word."
        }

        # Crear directorio de salida si no existe
        output_dir = Path(settings.GENERATED_DIR)
        output_dir.mkdir(parents=True, exist_ok=True)

        # Ruta de salida
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_salida = output_dir / f"test_cotizacion_{timestamp}.docx"

        print(f"\n📝 Generando cotización...")
        print(f"   Cliente: {datos_cotizacion['cliente']}")
        print(f"   Proyecto: {datos_cotizacion['proyecto']}")
        print(f"   Items: {len(datos_cotizacion['items'])}")
        print(f"   Total: S/ {datos_cotizacion['total']:,.2f}")
        print(f"   Archivo: {ruta_salida.name}")

        # Generar documento
        resultado = word_generator.generar_cotizacion(
            datos=datos_cotizacion,
            ruta_salida=str(ruta_salida),
            opciones={
                "mostrarPreciosUnitarios": True,
                "mostrarPreciosTotales": True,
                "mostrarIGV": True,
                "incluirLogo": False
            }
        )

        # Verificar resultado
        if ruta_salida.exists():
            file_size = ruta_salida.stat().st_size
            print(f"\n✅ ÉXITO: Archivo generado correctamente")
            print(f"   Ruta: {ruta_salida}")
            print(f"   Tamaño: {file_size:,} bytes ({file_size/1024:.2f} KB)")

            # Verificar que no está vacío
            if file_size > 0:
                print(f"   Estado: ✅ Archivo válido (no vacío)")

                # Intentar abrir con python-docx para verificar integridad
                try:
                    from docx import Document
                    doc = Document(str(ruta_salida))
                    num_paragraphs = len(doc.paragraphs)
                    num_tables = len(doc.tables)
                    print(f"   Párrafos: {num_paragraphs}")
                    print(f"   Tablas: {num_tables}")
                    print(f"   Integridad: ✅ Archivo no corrupto")
                except Exception as e:
                    print(f"   Integridad: ❌ Posible corrupción: {e}")
                    return False

                return True
            else:
                print(f"   Estado: ❌ Archivo vacío")
                return False
        else:
            print(f"\n❌ ERROR: Archivo no fue generado")
            print(f"   Resultado: {resultado}")
            return False

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def test_generar_proyecto_word():
    """Test 2: Generar informe de proyecto en Word"""
    print_section("TEST 2: Generación de Informe de Proyecto Word")

    try:
        # Datos de prueba para proyecto
        datos_proyecto = {
            "nombre": "PROYECTO DE PRUEBA - Instalación Eléctrica Edificio",
            "cliente": "CONSTRUCTORA DE PRUEBA S.A.C.",
            "descripcion": "Proyecto de prueba para verificar generación de informes",
            "presupuesto_estimado": 150000.00,
            "duracion_meses": 6,
            "fecha_inicio": datetime.now().strftime("%d/%m/%Y"),
            "estado": "En planificación",
            "fases": [
                {
                    "nombre": "Fase 1: Planificación",
                    "duracion": "1 mes",
                    "actividades": ["Revisión de planos", "Cotización de materiales"]
                },
                {
                    "nombre": "Fase 2: Instalación",
                    "duracion": "4 meses",
                    "actividades": ["Instalación de tableros", "Cableado", "Pruebas"]
                },
                {
                    "nombre": "Fase 3: Entrega",
                    "duracion": "1 mes",
                    "actividades": ["Pruebas finales", "Documentación", "Capacitación"]
                }
            ],
            "recursos": [
                {"tipo": "Personal", "descripcion": "Ingeniero Eléctrico", "cantidad": 1},
                {"tipo": "Personal", "descripcion": "Técnicos electricistas", "cantidad": 4},
                {"tipo": "Material", "descripcion": "Tableros eléctricos", "cantidad": 15},
                {"tipo": "Material", "descripcion": "Cable NYY (metros)", "cantidad": 5000}
            ]
        }

        # Crear directorio de salida
        output_dir = Path(settings.GENERATED_DIR)
        output_dir.mkdir(parents=True, exist_ok=True)

        # Ruta de salida
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_salida = output_dir / f"test_proyecto_{timestamp}.docx"

        print(f"\n📊 Generando informe de proyecto...")
        print(f"   Cliente: {datos_proyecto['cliente']}")
        print(f"   Proyecto: {datos_proyecto['nombre']}")
        print(f"   Presupuesto: S/ {datos_proyecto['presupuesto_estimado']:,.2f}")
        print(f"   Duración: {datos_proyecto['duracion_meses']} meses")
        print(f"   Archivo: {ruta_salida.name}")

        # Generar documento
        resultado = word_generator.generar_informe_proyecto(
            datos=datos_proyecto,
            ruta_salida=str(ruta_salida),
            opciones={
                "incluir_cronograma": True,
                "incluir_recursos": True,
                "incluir_analisis": True
            }
        )

        # Verificar resultado
        if ruta_salida.exists():
            file_size = ruta_salida.stat().st_size
            print(f"\n✅ ÉXITO: Archivo generado correctamente")
            print(f"   Ruta: {ruta_salida}")
            print(f"   Tamaño: {file_size:,} bytes ({file_size/1024:.2f} KB)")

            if file_size > 0:
                print(f"   Estado: ✅ Archivo válido (no vacío)")

                # Verificar integridad
                try:
                    from docx import Document
                    doc = Document(str(ruta_salida))
                    num_paragraphs = len(doc.paragraphs)
                    num_tables = len(doc.tables)
                    print(f"   Párrafos: {num_paragraphs}")
                    print(f"   Tablas: {num_tables}")
                    print(f"   Integridad: ✅ Archivo no corrupto")
                except Exception as e:
                    print(f"   Integridad: ❌ Posible corrupción: {e}")
                    return False

                return True
            else:
                print(f"   Estado: ❌ Archivo vacío")
                return False
        else:
            print(f"\n❌ ERROR: Archivo no fue generado")
            print(f"   Resultado: {resultado}")
            return False

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Ejecutar todas las pruebas"""
    print("\n" + "="*70)
    print("  🧪 PRUEBAS DE GENERACIÓN DE DOCUMENTOS WORD")
    print("  TESLA COTIZADOR V3.0")
    print("="*70)

    # Verificar que word_generator está disponible
    if word_generator is None:
        print("\n❌ ERROR CRÍTICO: word_generator no está inicializado")
        return False

    print(f"\n✅ WordGenerator inicializado correctamente")
    print(f"   Directorio de salida: {settings.GENERATED_DIR}")

    # Ejecutar pruebas
    resultados = []

    # Test 1: Cotización
    resultado_cotizacion = test_generar_cotizacion_word()
    resultados.append(("Cotización Word", resultado_cotizacion))

    # Test 2: Proyecto
    resultado_proyecto = test_generar_proyecto_word()
    resultados.append(("Proyecto Word", resultado_proyecto))

    # Resumen
    print_section("RESUMEN DE PRUEBAS")

    for nombre, resultado in resultados:
        estado = "✅ PASS" if resultado else "❌ FAIL"
        print(f"  {estado}  {nombre}")

    total_pass = sum(1 for _, r in resultados if r)
    total = len(resultados)

    print(f"\n  Total: {total_pass}/{total} pruebas pasadas")

    if total_pass == total:
        print(f"\n🎉 TODAS LAS PRUEBAS PASARON CORRECTAMENTE")
        return True
    else:
        print(f"\n⚠️  ALGUNAS PRUEBAS FALLARON")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
